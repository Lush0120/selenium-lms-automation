"""
services/document_generator.py
Generador de documentos Word para correos de asignación de pruebas.

Genera documentos completos con:
- Saludo personalizado
- Tabla dinámica de pruebas con niveles
- Credenciales de acceso
- Fechas de habilitación
- Instrucciones y firma

Uso:
    from services.document_generator import DocumentGenerator
    
    generator = DocumentGenerator()
    
    # Generar documento
    filepath = generator.generate(
        nombre_candidato="Juan García",
        email="juan@example.com",
        usuario="jgarcia",
        contrasena="Abc123!",
        pruebas=lista_pruebas,
        duracion_dias=2
    )
"""

import os
import base64
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.logger import get_logger

logger = get_logger(__name__)


@dataclass
class PruebaInfo:
    """Información de una prueba para el documento."""
    nombre: str
    url: str
    niveles: List[Dict[str, str]]  # [{"nivel": "Básico", "tiempo": "15 min"}, ...]
    lenguaje: str = "N/A"
    monitorizacion: str = "N/A"


class DocumentGenerator:
    """Generador de documentos Word para asignación de pruebas."""
    
    # Colores
    COLOR_HEADER = RGBColor(47, 84, 105)  # Azul oscuro para encabezados
    COLOR_WHITE = RGBColor(255, 255, 255)
    
    # Meses en español
    MESES = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }
    
    def __init__(self, output_dir: str = "."):
        """
        Inicializa el generador.
        
        Args:
            output_dir: Directorio donde guardar los documentos generados
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Buscar imagen de firma
        self.firma_path = self._find_firma_image()
    
    def _find_firma_image(self) -> Optional[Path]:
        """Busca la imagen de firma en ubicaciones comunes."""
        possible_paths = [
            Path(__file__).parent.parent / "assets" / "frima_qvision.jfif",
            Path(__file__).parent.parent / "firma_qvision.jfif",
            Path("firma_qvision.jfif"),
            Path("assets/firma_qvision.jfif"),
        ]
        
        for path in possible_paths:
            if path.exists():
                return path
        
        logger.warning("No se encontró imagen de firma")
        return None
    
    def generate(
        self,
        nombre_candidato: str,
        email: str,
        usuario: str,
        contrasena: str,
        pruebas: List[PruebaInfo],
        duracion_dias: int = 1,
        es_usuario_nuevo: bool = False,
        fecha_inicio: Optional[datetime] = None
    ) -> str:
        """
        Genera el documento Word completo.
        
        Args:
            nombre_candidato: Nombre completo del candidato
            email: Email del candidato
            usuario: Username de Moodle
            contrasena: Contraseña generada
            pruebas: Lista de PruebaInfo con datos de las pruebas
            duracion_dias: Días de habilitación (1, 2, 3)
            es_usuario_nuevo: Si es usuario nuevo o existente
            fecha_inicio: Fecha de inicio (default: ahora)
            
        Returns:
            Ruta del archivo generado
        """
        logger.info(f"Generando documento para: {nombre_candidato}")
        
        # Calcular fechas
        if fecha_inicio is None:
            fecha_inicio = datetime.now()
        fecha_fin = fecha_inicio + timedelta(days=duracion_dias)
        
        # Crear documento
        doc = Document()
        
        # Configurar estilos
        self._setup_styles(doc)
        
        # Agregar contenido
        self._add_saludo(doc, nombre_candidato)
        self._add_intro(doc)
        self._add_tabla_pruebas(doc, pruebas)
        self._add_consideraciones(doc)
        self._add_credenciales(doc, usuario, contrasena, es_usuario_nuevo)
        self._add_fechas(doc, fecha_inicio, fecha_fin, duracion_dias)
        self._add_notas(doc)
        self._add_instrucciones_smowl(doc)
        self._add_despedida(doc)
        self._add_firma(doc)
        self._add_disclaimer(doc)
        
        # Guardar documento
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"Asignacion_Pruebas_{nombre_candidato.replace(' ', '_')}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        doc.save(filepath)
        logger.info(f"Documento generado: {filepath}")
        
        return str(filepath)
    
    def _setup_styles(self, doc: Document):
        """Configura los estilos del documento."""
        # Estilo normal
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Configurar fuente para caracteres especiales
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    
    def _add_saludo(self, doc: Document, nombre: str):
        """Agrega el saludo inicial."""
        p = doc.add_paragraph()
        p.add_run(f"Hola {nombre},").bold = False
        doc.add_paragraph()  # Línea vacía
    
    def _add_intro(self, doc: Document):
        """Agrega el párrafo introductorio."""
        intro = (
            "Q-Vision Technologies le da un caluroso saludo y le informa que, "
            "para continuar con el proceso de selección en el que usted se encuentra "
            "participando, es necesario que disponga del tiempo suficiente para "
            "realizar la(s) prueba(s) cuyos detalles se encuentra relacionados en la "
            "siguiente tabla:"
        )
        doc.add_paragraph(intro)
        doc.add_paragraph()  # Línea vacía
    
    def _add_tabla_pruebas(self, doc: Document, pruebas: List[PruebaInfo]):
        """
        Agrega la tabla de pruebas con niveles dinámicos.
        
        La tabla maneja celdas combinadas (merge) cuando hay múltiples niveles.
        """
        # Calcular total de filas necesarias
        total_rows = 1  # Header
        for prueba in pruebas:
            total_rows += max(1, len(prueba.niveles))
        
        # Crear tabla con 6 columnas
        # Nombre | Link | Nivel | Tiempo | Lenguaje | Monitorización
        table = doc.add_table(rows=total_rows, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar anchos de columnas (en pulgadas)
        widths = [1.8, 2.5, 1.0, 0.8, 0.8, 0.9]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
        
        # Header
        header_texts = [
            "Nombre de la prueba",
            "Link de acceso a la prueba", 
            "Nivel de conocimiento a evaluar",
            "Tiempo disponible por nivel",
            "Lenguaje",
            "Monitorización"
        ]
        
        header_row = table.rows[0]
        for i, text in enumerate(header_texts):
            cell = header_row.cells[i]
            cell.text = text
            self._format_header_cell(cell)
        
        # Datos de pruebas
        current_row = 1
        for prueba in pruebas:
            num_niveles = max(1, len(prueba.niveles))
            
            # Primera fila de la prueba
            row = table.rows[current_row]
            
            # Nombre (se combina si hay múltiples niveles)
            cell_nombre = row.cells[0]
            cell_nombre.text = prueba.nombre
            self._format_data_cell(cell_nombre)
            
            # Link (se combina si hay múltiples niveles)
            cell_link = row.cells[1]
            # Agregar como hipervínculo
            self._add_hyperlink(cell_link, prueba.url, prueba.url)
            self._format_data_cell(cell_link, is_link=True)
            
            # Lenguaje (solo en primera fila, se combina si hay múltiples)
            cell_lenguaje = row.cells[4]
            cell_lenguaje.text = prueba.lenguaje
            self._format_data_cell(cell_lenguaje)
            
            # Monitorización (solo en primera fila, se combina si hay múltiples)
            cell_monit = row.cells[5]
            cell_monit.text = prueba.monitorizacion
            self._format_data_cell(cell_monit)
            
            # Niveles y tiempos
            if prueba.niveles:
                for i, nivel_info in enumerate(prueba.niveles):
                    if i == 0:
                        # Primera fila ya creada
                        row.cells[2].text = nivel_info.get("nivel", "N/A")
                        row.cells[3].text = nivel_info.get("tiempo", "N/A")
                        self._format_data_cell(row.cells[2])
                        self._format_data_cell(row.cells[3])
                    else:
                        # Filas adicionales para más niveles
                        next_row = table.rows[current_row + i]
                        next_row.cells[2].text = nivel_info.get("nivel", "N/A")
                        next_row.cells[3].text = nivel_info.get("tiempo", "N/A")
                        self._format_data_cell(next_row.cells[2])
                        self._format_data_cell(next_row.cells[3])
            else:
                # Sin niveles específicos
                row.cells[2].text = "N/A"
                row.cells[3].text = "N/A"
                self._format_data_cell(row.cells[2])
                self._format_data_cell(row.cells[3])
            
            # Combinar celdas verticalmente si hay múltiples niveles
            if num_niveles > 1:
                # Combinar Nombre
                self._merge_cells_vertical(table, current_row, current_row + num_niveles - 1, 0)
                # Combinar Link
                self._merge_cells_vertical(table, current_row, current_row + num_niveles - 1, 1)
                # Combinar Lenguaje
                self._merge_cells_vertical(table, current_row, current_row + num_niveles - 1, 4)
                # Combinar Monitorización
                self._merge_cells_vertical(table, current_row, current_row + num_niveles - 1, 5)
            
            current_row += num_niveles
        
        doc.add_paragraph()  # Línea vacía después de la tabla
    
    def _merge_cells_vertical(self, table, start_row: int, end_row: int, col: int):
        """Combina celdas verticalmente."""
        if start_row >= end_row:
            return
        
        start_cell = table.rows[start_row].cells[col]
        for row_idx in range(start_row + 1, end_row + 1):
            cell = table.rows[row_idx].cells[col]
            start_cell.merge(cell)
        
        # Centrar verticalmente
        start_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def _format_header_cell(self, cell):
        """Formatea una celda de encabezado."""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Color de fondo
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2F5469')  # Azul oscuro
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Texto en blanco y negrita
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = self.COLOR_WHITE
                run.font.bold = True
                run.font.size = Pt(10)
    
    def _format_data_cell(self, cell, is_link: bool = False):
        """Formatea una celda de datos."""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if is_link:
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
    
    def _add_hyperlink(self, cell, url: str, text: str):
        """Agrega un hipervínculo a una celda."""
        # Limpiar celda
        cell.text = ""
        
        paragraph = cell.paragraphs[0]
        
        # Crear el hipervínculo
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Color azul
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        # Subrayado
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        # Tamaño de fuente
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')  # 10pt
        rPr.append(sz)
        
        new_run.append(rPr)
        
        # Texto del enlace
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _add_consideraciones(self, doc: Document):
        """Agrega la sección de consideraciones."""
        p = doc.add_paragraph()
        p.add_run("Consideraciones para la presentación de la prueba:").bold = True
        doc.add_paragraph()
    
    def _add_credenciales(self, doc: Document, usuario: str, contrasena: str, es_nuevo: bool):
        """Agrega las credenciales de acceso."""
        if es_nuevo:
            texto = "Se ha creado una cuenta para usted en la plataforma. Por favor, acceda a las pruebas con las siguientes credenciales:"
        else:
            texto = "Actualmente usted es usuario(a) activo(a) en la plataforma, por favor, acceda a las pruebas con sus credenciales de acceso:"
        
        doc.add_paragraph(texto)
        doc.add_paragraph()
        
        p1 = doc.add_paragraph()
        p1.add_run("Usuario: ").bold = True
        p1.add_run(usuario)
        
        p2 = doc.add_paragraph()
        p2.add_run("Contraseña: ").bold = True
        p2.add_run(contrasena)
        
        doc.add_paragraph()
    
    def _add_fechas(self, doc: Document, inicio: datetime, fin: datetime, dias: int):
        """Agrega las fechas de habilitación."""
        p = doc.add_paragraph()
        p.add_run(f"1. La(s) prueba(s) se encuentra(n) activa(s) durante {dias * 24} horas:").bold = False
        
        doc.add_paragraph()
        
        p_inicio = doc.add_paragraph()
        p_inicio.add_run("Fecha y Hora inicio: ").bold = True
        p_inicio.add_run(self._format_fecha(inicio))
        
        p_fin = doc.add_paragraph()
        p_fin.add_run("Fecha y Hora Finalización: ").bold = True
        p_fin.add_run(self._format_fecha(fin))
        
        doc.add_paragraph()
        
        doc.add_paragraph("2. Una vez terminada la prueba, es necesario dar respuesta a este correo informando que ya fue realizada.")
        doc.add_paragraph()
    
    def _format_fecha(self, fecha: datetime) -> str:
        """Formatea una fecha en español."""
        mes = self.MESES[fecha.month]
        return f"{fecha.day} de {mes} de {fecha.year}, {fecha.strftime('%I:%M %p')}"
    
    def _add_notas(self, doc: Document):
        """Agrega las notas importantes."""
        p = doc.add_paragraph()
        p.add_run("NOTAS:").bold = True
        
        doc.add_paragraph()
        
        notas = [
            "Evite buscar respuestas en línea, porque la plataforma registrará las ocasiones y el tiempo que salió de la ventana activa para la prueba en curso.",
            "En caso de tener cualquier inconveniente con la prueba, contacte a la generalista de People que está a cargo de su proceso de selección.",
            "En caso de ser asignado a pruebas Evalart, las instrucciones para realizarlas y el link de acceso le serán enviados en un correo adicional."
        ]
        
        for i, nota in enumerate(notas, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {nota}")
        
        doc.add_paragraph()
    
    def _add_instrucciones_smowl(self, doc: Document):
        """Agrega las instrucciones de SMOWL."""
        p = doc.add_paragraph()
        p.add_run("🖥️ Instalación de la extensión de monitorización").bold = True
        
        doc.add_paragraph()
        
        p1 = doc.add_paragraph()
        p1.add_run("1. Revisa el siguiente ")
        self._add_hyperlink_paragraph(p1, "https://vimeo.com/1123715258/6dbd0b736e?share=copy&fl=sv&fe=ci", "video tutorial")
        p1.add_run(" para instalar la extensión ")
        run_smowl = p1.add_run("SMOWL")
        run_smowl.bold = True
        p1.add_run(" en tu computador.")
        
        doc.add_paragraph("2. Acepta los términos y condiciones explicados en el video anterior.")
        doc.add_paragraph()
    
    def _add_hyperlink_paragraph(self, paragraph, url: str, text: str):
        """Agrega un hipervínculo a un párrafo existente."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        b = OxmlElement('w:b')
        rPr.append(b)
        
        new_run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _add_despedida(self, doc: Document):
        """Agrega la despedida."""
        doc.add_paragraph("Estamos atentos para apoyarte en lo que necesites.")
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.add_run("¡Éxitos en tu prueba!").bold = True
        
        doc.add_paragraph()
        doc.add_paragraph("Cordialmente,")
        doc.add_paragraph()
    
    def _add_firma(self, doc: Document):
        """Agrega la imagen de firma."""
        if self.firma_path and self.firma_path.exists():
            try:
                doc.add_picture(str(self.firma_path), width=Inches(4.5))
            except Exception as e:
                logger.warning(f"No se pudo agregar imagen de firma: {e}")
                doc.add_paragraph("[Firma Q-Vision Technologies]")
        else:
            doc.add_paragraph("[Firma Q-Vision Technologies]")
        
        doc.add_paragraph()
    
    def _add_disclaimer(self, doc: Document):
        """Agrega el disclaimer final."""
        disclaimer = (
            '"Si recibes este correo por fuera de tu jornada laboral, puedes darle '
            'respuesta cuando retomes tu horario habitual, salvo causa de fuerza '
            'mayor o por circunstancias excepcionales establecidas en la política de '
            'desconexión laboral."'
        )
        
        p = doc.add_paragraph()
        run = p.add_run(disclaimer)
        run.italic = True
        run.font.size = Pt(9)
    
    def generate_base64(
        self,
        nombre_candidato: str,
        email: str,
        usuario: str,
        contrasena: str,
        pruebas: List[PruebaInfo],
        duracion_dias: int = 1,
        es_usuario_nuevo: bool = False,
        fecha_inicio: Optional[datetime] = None
    ) -> tuple[str, str]:
        """
        Genera el documento y retorna como base64 para envío por API.
        
        Returns:
            Tupla (nombre_archivo, contenido_base64)
        """
        filepath = self.generate(
            nombre_candidato=nombre_candidato,
            email=email,
            usuario=usuario,
            contrasena=contrasena,
            pruebas=pruebas,
            duracion_dias=duracion_dias,
            es_usuario_nuevo=es_usuario_nuevo,
            fecha_inicio=fecha_inicio
        )
        
        with open(filepath, 'rb') as f:
            content_base64 = base64.b64encode(f.read()).decode('utf-8')
        
        filename = Path(filepath).name
        return filename, content_base64


def pruebas_from_course_data(courses: List[Any]) -> List[PruebaInfo]:
    """
    Convierte una lista de CourseData a PruebaInfo para el generador.
    
    Usa el método get_table_data() de CourseData si está disponible,
    sino hace la extracción manual.
    
    Args:
        courses: Lista de CourseData del cache
        
    Returns:
        Lista de PruebaInfo listos para el documento
    """
    pruebas = []
    
    for course in courses:
        # Intentar usar get_table_data() si está disponible
        if hasattr(course, 'get_table_data'):
            table_data = course.get_table_data()
            prueba = PruebaInfo(
                nombre=table_data.nombre,
                url=table_data.url,
                niveles=table_data.niveles,
                lenguaje=table_data.lenguaje,
                monitorizacion=table_data.monitorizacion
            )
            pruebas.append(prueba)
            continue
        
        # Fallback: extracción manual para compatibilidad
        characteristics = {}
        if hasattr(course, 'characteristics'):
            for char in course.characteristics:
                characteristics[char.name.lower()] = char.value
        
        # Extraer niveles (características que empiezan con "Nivel")
        niveles = []
        duracion = None
        
        for char in getattr(course, 'characteristics', []):
            name_lower = char.name.lower()
            
            # Buscar niveles explícitos (Nivel Básico, Nivel Intermedio, etc.)
            if name_lower.startswith('nivel') and 'pregunta' in char.value.lower():
                nivel_nombre = char.name.replace('Nivel', '').replace('nivel', '').strip()
                if nivel_nombre:
                    niveles.append({
                        "nivel": nivel_nombre.capitalize(),
                        "tiempo": ""
                    })
            
            # Nivel de la prueba (formato técnico)
            elif name_lower == 'nivel de la prueba':
                niveles.append({
                    "nivel": char.value,
                    "tiempo": ""
                })
            
            # Buscar duración
            if any(kw in name_lower for kw in ['duración', 'duracion', 'disponibilidad']):
                duracion = char.value
        
        # Si encontramos niveles, asignar el tiempo a cada uno
        if niveles and duracion:
            for nivel in niveles:
                nivel["tiempo"] = duracion
        elif not niveles:
            # Si no hay niveles específicos, crear uno genérico
            nivel_value = characteristics.get('nivel de la prueba', 'General')
            tiempo_value = duracion or characteristics.get('disponibilidad prueba', 'N/A')
            niveles = [{"nivel": nivel_value, "tiempo": tiempo_value}]
        
        # Extraer lenguaje y monitorización
        lenguaje = (
            characteristics.get('lenguaje') or 
            characteristics.get('idioma') or 
            'N/A'
        )
        monitorizacion = (
            characteristics.get('monitorización') or
            characteristics.get('monitorizado') or
            characteristics.get('supervisión en línea') or
            'N/A'
        )
        
        # Obtener URL
        url = getattr(course, 'url_course', '') or ''
        
        prueba = PruebaInfo(
            nombre=course.name,
            url=url,
            niveles=niveles,
            lenguaje=lenguaje,
            monitorizacion=monitorizacion
        )
        pruebas.append(prueba)
    
    return pruebas